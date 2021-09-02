from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor


class WordFile():
	def __init__(self, resume):
		super(WordFile, self).__init__()
		self.presentation = resume.presentation
		self.abstract = resume.abstract
		self.identification = resume.identification
		self.address = resume.address
		self.academic_titles = resume.academic_titles
		self.complementary_courses = resume.complementary_courses
		self.professional_activities_list = resume.professional_activities_list
		self.lines_of_research = resume.lines_of_research
		self.projects_dict = resume.projects_dict
		self.other_professional_activities_dict = resume.other_professional_activities_dict
		self.areas_of_expertise = resume.areas_of_expertise
		self.languages = resume.languages
		self.awards = resume.awards
		self.bibliographic_productions_dict = resume.bibliographic_productions_dict
		self.technical_productions_dict = resume.technical_productions_dict

		self.document = Document()

		self.define_style()
		self.document.add_heading("Apresentação", 0) # Add a section
		self.add_presentation()
		self.add_abstract()
		self.add_identification()
		self.add_address()
		self.add_academic_titles()
		self.add_complementary_courses()
		self.document.add_heading("Atuação Profissional", 0) # Add a section
		self.add_professional_activities()
		self.document.add_heading("Linhas de pesquisa", 0) # Add a section
		self.add_lines_of_research()
		self.add_projects()
		self.add_other_professional_activities()
		self.document.add_heading("Áreas de atuação", 0) # Add a section
		self.add_numbered_table(self.areas_of_expertise)
		self.document.add_heading("Idiomas", 0) # Add a section
		self.add_languages()
		self.document.add_heading("Prêmios e títulos", 0) # Add a section
		self.add_awards()
		self.document.add_heading("Produções", 0) # Add a section
		self.add_productions(self.bibliographic_productions_dict, "Produção bibliográfica")
		self.add_productions(self.technical_productions_dict, "Produção técnica")

		# # self.add_incomplete_articles()

	def define_style(self):
		style = self.document.styles['Normal']
		font = style.font
		font.name = 'Arial'
		font.size = Pt(10)

	def add_presentation(self):
		self.document.add_heading(self.presentation[0], 1) # Add the name as a title

		for item in self.presentation[1:]: # Add the other items
			paragraph = self.document.add_paragraph(item)
			# Format paragraph
			paragraph_format = paragraph.paragraph_format
			paragraph_format.space_before = Pt(4)
			paragraph_format.space_after = Pt(4)

	def add_abstract(self):
		self.document.add_heading("Resumo", 1) # Add "Resumo" as a title
		
		self.abstract = self.abstract[0].upper() + self.abstract[1:] # First letter uppercased
		paragraph = self.document.add_paragraph(self.abstract)

		# Format paragraph
		paragraph_format = paragraph.paragraph_format
		paragraph_format.line_spacing = Pt(15)
	
		paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

	def add_identification(self):
		self.document.add_heading("Identificação", 1) # Add "Identificação" as a title

		table = self.document.add_table(rows=0, cols=2)

		for key, value in self.identification.items():
			row_cells = table.add_row().cells
			row_cells[0].text = key
			row_cells[1].text = value

	def add_address(self):
		self.document.add_heading("Endereço", 1) # Add "Endereço" as a title

		table = self.document.add_table(rows=0, cols=2)

		row_cells = table.add_row().cells
		row_cells[0].text = "Endereço Profissional"
		row_cells[1].text = self.address[0]
		
		for value in self.address[1:]:
			row_cells = table.add_row().cells
			row_cells[1].text = value

	def add_academic_titles(self):
		self.document.add_heading("Formação acadêmica/titulação", 1) # Add "Formação acadêmica/titulação" as a title
		table = self.document.add_table(rows=0, cols=2) # Create table

		for pos, academic_title in enumerate(self.academic_titles['academic_title']):
			row_cells = table.add_row().cells # Get cells from row
			row_cells[0].width = Pt(80) # Make the first cell smaller
			paragraph = row_cells[0].paragraphs[0] # Get the paragraph
			paragraph.add_run(self.academic_titles['year_range'][pos]).bold = True # Add a number for each research and make it bold
			run = paragraph.runs[0]
			font = run.font
			font.color.rgb = RGBColor.from_string('0b306b')
			
			row_cells[1].width = Pt(480) # Make the second cell bigger

			academic_title_paragraph = row_cells[1].paragraphs[0] # Get the cell first paragraph
			academic_title_paragraph.text = academic_title # Add the academic title to the first paragraph

			institution_paragraph = row_cells[1].add_paragraph() # Add a second paragraph
			institution_paragraph.text = self.academic_titles['institution'][pos] # Add the institution to the paragraph
			
			if self.academic_titles['project_title'][pos] != "": # If it has keywords
				project_title_paragraph = row_cells[1].add_paragraph() # Add the project title paragraph
				project_title_paragraph.text = self.academic_titles['project_title'][pos] # Add the project title content

			if self.academic_titles['advisor'][pos] != "": # If it has keywords
				advisor_paragraph = row_cells[1].add_paragraph() # Add the advisor paragraph
				advisor_paragraph.text = self.academic_titles['advisor'][pos] # Add the advisor content

			if self.academic_titles['scholarship'][pos] != "": # If it has keywords
				scholarship_paragraph = row_cells[1].add_paragraph() # Add the scholarship paragraph
				scholarship_paragraph.text = self.academic_titles['scholarship'][pos] # Add the scholarship content

			if self.academic_titles['key_words'][pos] != "": # If it has keywords
				keywords_paragraph = row_cells[1].add_paragraph() # Add the keywords paragraph
				keywords_paragraph.text = self.academic_titles['key_words'][pos] # Add the keywords content

	def add_complementary_courses(self):
		self.document.add_heading("Formação Complementar", 1) # Add "Formação acadêmica/titulação" as a title
		table = self.document.add_table(rows=0, cols=2) # Create table

		for pos, course in enumerate(self.complementary_courses['course_name']):
			row_cells = table.add_row().cells # Get cells from row
			row_cells[0].width = Pt(80) # Make the first cell smaller
			paragraph = row_cells[0].paragraphs[0] # Get the paragraph
			paragraph.add_run(self.complementary_courses['year_range'][pos]).bold = True # Add a number for each research and make it bold
			run = paragraph.runs[0]
			font = run.font
			font.color.rgb = RGBColor.from_string('0b306b')
			
			row_cells[1].width = Pt(480) # Make the second cell bigger

			course_paragraph = row_cells[1].paragraphs[0] # Get the cell first paragraph
			course_paragraph.text = course # Add the academic title to the first paragraph

			institution_paragraph = row_cells[1].add_paragraph() # Add a second paragraph
			institution_paragraph.text = self.complementary_courses['institution'][pos] # Add the institution to the paragraph

	def set_cell_background(self, cell, fill, color=None, val=None):
		from docx.oxml.shared import qn  # feel free to move these out
		from docx.oxml.xmlchemy import OxmlElement

		cell_properties = cell._element.tcPr
		try:
			cell_shading = cell_properties.xpath('w:shd')[0]  # in case there's already shading
		except IndexError:
			cell_shading = OxmlElement('w:shd') # add new w:shd element to it
		if fill:
			cell_shading.set(qn('w:fill'), fill)  # set fill property, respecting namespace
		if color:
			pass # TODO
		if val:
			pass # TODO
		cell_properties.append(cell_shading)  # finally extend cell props with shading element

	def add_subsection(self, subsection):
		table = self.document.add_table(rows=0, cols=1) # Create table row
		row_cells = table.add_row().cells # Get cells from first row
		paragraph = row_cells[0].paragraphs[0] # Get the paragraph

		paragraph_format = paragraph.paragraph_format # Get the paragraph format
		# Adjust space before and after
		paragraph_format.space_before = Pt(3)
		paragraph_format.space_after = Pt(3)

		# Add a new run to the paragraph, make the text bold and white and change the size and the background color
		paragraph.add_run(subsection).bold = True # Add a number for each publication and make it bold
		run = paragraph.runs[0]
		font = run.font
		font.size = Pt(13)
		font.color.rgb = RGBColor.from_string('FFFFFF')
		self.set_cell_background(row_cells[0], '0b306b')

	def add_professional_activities(self):
		for item in self.professional_activities_list:
			self.add_subsection(item['institution'])
			bonds_paragraph = self.document.add_paragraph()
			bonds_paragraph.add_run('Vínculo institucional').bold = True # Add the year range for each bond and make it bold
			run = bonds_paragraph.runs[0]
			font = run.font
			font.size = Pt(12)
			
			# Format paragraph
			bonds_paragraph_format = bonds_paragraph.paragraph_format
			bonds_paragraph_format.space_before = Pt(4)
			bonds_paragraph_format.space_after = Pt(4)

			table = self.document.add_table(rows=0, cols=2) # Create table

			for bond in item['Vínculo institucional']:
				row_cells = table.add_row().cells # Get cells from row
				row_cells[0].width = Pt(120) # Make the first cell smaller
				paragraph = row_cells[0].paragraphs[0] # Get the paragraph
				paragraph.add_run(bond['year_range']).bold = True # Add the year range for each bond and make it bold
				run = paragraph.runs[0]
				font = run.font
				font.color.rgb = RGBColor.from_string('0b306b')
				
				row_cells[1].width = Pt(440) # Make the second cell bigger

				bond_paragraph = row_cells[1].paragraphs[0] # Get the cell first paragraph
				bond_paragraph.text = bond['content'] # Add the bond content to the paragraph
				if len(bond['content']) >= 70:
					bond_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

				if 'second_line' in bond.keys():
					row_cells = table.add_row().cells # Get cells from row
					row_cells[0].width = Pt(120) # Make the first cell smaller
					paragraph = row_cells[0].paragraphs[0] # Get the paragraph
					paragraph.add_run(bond['second_line']).bold = True # Add the year range for each bond and make it bold
					run = paragraph.runs[0]
					font = run.font
					font.color.rgb = RGBColor.from_string('0b306b')
					
					row_cells[1].width = Pt(440) # Make the second cell bigger

					second_line_paragraph = row_cells[1].paragraphs[0] # Get the cell first paragraph
					second_line_paragraph.text = bond['second_line_content'] # Add the bond content to the paragraph
					if len(bond['second_line_content']) >= 70:
						second_line_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

			if item['Atividades'] is not None:
				activities_paragraph = self.document.add_paragraph()
				activities_paragraph.add_run('Atividades').bold = True # Add the year range for each bond and make it bold
				run = activities_paragraph.runs[0]
				font = run.font
				font.size = Pt(12)

				activities_paragraph_format = activities_paragraph.paragraph_format
				activities_paragraph_format.space_before = Pt(4)
				activities_paragraph_format.space_after = Pt(4)

				table2 = self.document.add_table(rows=0, cols=2) # Create table
				
				for activity in item['Atividades']:
					row_cells2 = table2.add_row().cells # Get cells from row
					row_cells2[0].width = Pt(120) # Make the first cell smaller
					paragraph2 = row_cells2[0].paragraphs[0] # Get the paragraph
					paragraph2.add_run(activity['year_range']).bold = True # Add the year range for each activity and make it bold
					run2 = paragraph2.runs[0]
					font2 = run2.font
					font2.color.rgb = RGBColor.from_string('0b306b')
				
					row_cells2[1].width = Pt(440) # Make the second cell bigger

					activity_paragraph = row_cells2[1].paragraphs[0] # Get the cell first paragraph
					activity_paragraph.text = activity['content'] # Add the activity content to the paragraph

	def add_lines_of_research(self):
		table = self.document.add_table(rows=0, cols=2) # Create table
		for pos, research in enumerate(self.lines_of_research['title']):
			row_cells = table.add_row().cells # Get cells from row
			row_cells[0].width = 5 # Make the first cell smaller
			paragraph = row_cells[0].paragraphs[0] # Get the paragraph
			paragraph.add_run(str(pos + 1)).bold = True # Add a number for each research and make it bold
			run = paragraph.runs[0]
			font = run.font
			font.color.rgb = RGBColor.from_string('0b306b')
			
			row_cells[1].width = Pt(500) # Make the second cell bigger

			title_paragraph = row_cells[1].paragraphs[0] # Get the cell first paragraph
			title_paragraph.text = research # Add the title to the first paragraph

			goals_paragraph = row_cells[1].add_paragraph() # Add a second paragraph
			goals_paragraph.text = self.lines_of_research['goals'][pos] # Add the goals to the paragraph
			goals_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # Justify the paragraph

			if self.lines_of_research['key_words'][pos] != "": # If it has keywords
				keywords_paragraph = row_cells[1].add_paragraph() # Add the keywords paragraph
				keywords_paragraph.text = self.lines_of_research['key_words'][pos] # Add the keywords to the paragraph

	def add_projects(self):
		for project_nature in self.projects_dict.keys():
			self.document.add_heading(project_nature, 0) # Add a section
			
			projects_list = self.projects_dict[project_nature]
			table = self.document.add_table(rows=0, cols=2) # Create table

			for pos, research in enumerate(projects_list['title']):
				row_cells = table.add_row().cells # Get cells from row
				row_cells[0].width = Pt(80) # Make the first cell smaller
				paragraph = row_cells[0].paragraphs[0] # Get the paragraph
				paragraph.add_run(projects_list['year_range'][pos]).bold = True # Add a number for each research and make it bold
				run = paragraph.runs[0]
				font = run.font
				font.color.rgb = RGBColor.from_string('0b306b')
				
				row_cells[1].width = Pt(480) # Make the second cell bigger

				title_paragraph = row_cells[1].paragraphs[0] # Get the cell first paragraph
				title_paragraph.text = research # Add the title to the first paragraph

				description_paragraph = row_cells[1].add_paragraph() # Add a description paragraph
				description_paragraph.text = projects_list['description'][pos] # Add the description paragraph content
				description_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # Justify the paragraph

				# Format description paragraph
				paragraph_format = description_paragraph.paragraph_format
				paragraph_format.space_after = Pt(2)

				sit_nat_paragraph = row_cells[1].add_paragraph() # Add a situation/nature paragraph
				sit_nat_paragraph.text = projects_list['situation/nature'][pos] # Add the situation/nature paragraph content

				# Format situation/nature paragraph
				paragraph_format = sit_nat_paragraph.paragraph_format
				paragraph_format.space_before = Pt(0)
				paragraph_format.space_after = Pt(2)

				students_paragraph = row_cells[1].add_paragraph() # Add a students paragraph
				students_paragraph.text = projects_list['students'][pos] # Add the students paragraph content
				students_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # Justify the paragraph

				members_paragraph = row_cells[1].add_paragraph() # Add a members paragraph
				members_paragraph.text = projects_list['members'][pos] # Add the members paragraph content
				if len(projects_list['members'][pos]) >= 70:
					members_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # Justify the paragraph

				# Format members paragraph
				paragraph_format = members_paragraph.paragraph_format
				paragraph_format.space_after = Pt(2)

				if projects_list['financiers'][pos] is not None: # If there's financiers
					financiers_paragraph = row_cells[1].add_paragraph() # Add a financiers paragraph
					financiers_paragraph.text = projects_list['financiers'][pos] # Add the financiers paragraph content
					if len(projects_list['financiers'][pos]) >= 70:
						financiers_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # Justify the paragraph

					# Format financiers paragraph
					paragraph_format = financiers_paragraph.paragraph_format
					paragraph_format.space_before = Pt(0)
					paragraph_format.space_after = Pt(2)

				cta_paragraph = row_cells[1].add_paragraph() # Add a num_of_productions paragraph
				cta_paragraph.text = projects_list['num_of_productions'][pos] # Add the num_of_productions paragraph content
				
				if projects_list['financiers'][pos] is None: # If there's no financiers the cta paragraph will be the second one so it has a 0 space before
					paragraph_format = cta_paragraph.paragraph_format
					paragraph_format.space_before = Pt(0)

	def add_other_professional_activities(self):
		for activity_type in self.other_professional_activities_dict.keys():
			self.document.add_heading(activity_type, 0) # Add a section
			
			activities_list = self.other_professional_activities_dict[activity_type]
			table = self.document.add_table(rows=0, cols=2) # Create table

			for pos, activity in enumerate(activities_list['institution']):
				row_cells = table.add_row().cells # Get cells from row
				row_cells[0].width = Pt(80) # Make the first cell smaller
				paragraph = row_cells[0].paragraphs[0] # Get the paragraph
				paragraph.add_run(activities_list['year_range'][pos]).bold = True # Add a number for each activity and make it bold
				run = paragraph.runs[0]
				font = run.font
				font.color.rgb = RGBColor.from_string('0b306b')
				
				row_cells[1].width = Pt(480) # Make the second cell bigger

				title_paragraph = row_cells[1].paragraphs[0] # Get the cell first paragraph
				title_paragraph.text = activity # Add the title to the first paragraph
				if len(activity) >= 70:
					title_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # Justify the paragraph

	def add_languages(self):
		table = self.document.add_table(rows=0, cols=2) # Create table

		for language, string in zip(self.languages['language'], self.languages['string']):
			row_cells = table.add_row().cells # Get cells from row
			row_cells[0].width = 5 # Make the first cell smaller
			paragraph = row_cells[0].paragraphs[0] # Get the paragraph
			paragraph.add_run(language).bold = True # Add a number for each info and make it bold
			run = paragraph.runs[0]
			font = run.font
			font.color.rgb = RGBColor.from_string('0b306b')
			
			row_cells[1].width = Pt(500) # Make the second cell bigger
			paragraph = row_cells[1].paragraphs[0] # Get the cell paragraph
			paragraph.text = string # Add the info to the paragraph
			paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # Justify the paragraph

	def add_awards(self):
		table = self.document.add_table(rows=0, cols=2) # Create table

		for year, string in zip(self.awards['year'], self.awards['string']):
			row_cells = table.add_row().cells # Get cells from row
			row_cells[0].width = 5 # Make the first cell smaller
			paragraph = row_cells[0].paragraphs[0] # Get the paragraph
			paragraph.add_run(year).bold = True # Add a number for each info and make it bold
			run = paragraph.runs[0]
			font = run.font
			font.color.rgb = RGBColor.from_string('0b306b')
			
			row_cells[1].width = Pt(500) # Make the second cell bigger
			paragraph = row_cells[1].paragraphs[0] # Get the cell paragraph
			paragraph.text = string # Add the info to the paragraph
			paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # Justify the paragraph

	def add_numbered_table(self, info_list):
		table = self.document.add_table(rows=0, cols=2) # Create table

		for pos, info in enumerate(info_list):
			row_cells = table.add_row().cells # Get cells from row
			row_cells[0].width = 5 # Make the first cell smaller
			paragraph = row_cells[0].paragraphs[0] # Get the paragraph
			paragraph.add_run(str(pos + 1)).bold = True # Add a number for each info and make it bold
			run = paragraph.runs[0]
			font = run.font
			font.color.rgb = RGBColor.from_string('0b306b')
			
			row_cells[1].width = Pt(500) # Make the second cell bigger
			paragraph = row_cells[1].paragraphs[0] # Get the cell paragraph
			paragraph.text = info # Add the info to the paragraph
			paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # Justify the paragraph

	def add_productions(self, productions_dict, subsection):
		self.add_subsection(subsection)

		for key, publication_type in productions_dict.items():
			self.document.add_heading(key, 1) # Add the key as a title

			self.add_numbered_table(publication_type)

	def save_document(self, document_name):
		self.document.save(f'Files/{document_name}.docx')